"""
TranslateDocumentTool — translate document content to another language.

Zone: GREEN — runs automatically, no human approval required.

Supports:
    - PDF
    - DOCX
    - TXT
    - Extensionless files downloaded from Google Drive

Translation strategy:
    1. Google Cloud Translate, if configured and available.
    2. LLM translation fallback, if an LLM callback is available.
    3. deep-translator fallback, if installed.

The translated result is saved as a ready-to-use DOCX by default.
"""

from __future__ import annotations

import json
import logging
import mimetypes
import os
import re
import tempfile
from typing import Any, Callable

from core.tools.base_tool import BaseTool, ToolZone

logger = logging.getLogger(__name__)


_LANG_NAMES = {
    "en": "English",
    "ta": "Tamil",
    "hi": "Hindi",
    "ml": "Malayalam",
    "te": "Telugu",
    "kn": "Kannada",
    "fr": "French",
    "de": "German",
    "es": "Spanish",
    "ar": "Arabic",
    "zh": "Chinese",
    "ja": "Japanese",
    "pt": "Portuguese",
    "ru": "Russian",
}


class TranslateDocumentTool(BaseTool):
    """
    Translate the text content of a document to another language.

    Input:

        {
            "file_path": "/tmp/report.pdf",
            "target_lang": "ta",
            "source_lang": "en",
            "output_format": "docx"
        }

    Returns:

        {
            "status": "translated",
            "file_path": "/tmp/report_ta.docx",
            "filename": "report_ta.docx",
            "source_lang": "en",
            "target_lang": "ta",
            "target_name": "Tamil",
            "word_count": 420
        }
    """

    name: str = "translate_document"

    description: str = (
        "Translate a document (PDF/DOCX/TXT) to another language. "
        "GREEN — runs automatically. "
        "Input JSON: "
        "{\"file_path\": \"/tmp/report.pdf\", \"target_lang\": \"ta\"}. "
        "Language codes: ta=Tamil, hi=Hindi, fr=French, de=German, "
        "es=Spanish, ar=Arabic, zh=Chinese, ja=Japanese. "
        "The returned file_path is the complete translated Word document. "
        "Do NOT call generate_content or create_docx after this. "
        "If the user wants it saved to Drive, pass the returned file_path "
        "to upload_to_drive."
    )

    zone: ToolZone = ToolZone.GREEN

    def __init__(
        self,
        workspace_id: str | None = None,
        llm_translate_callback: Callable[..., str] | None = None,
    ) -> None:
        self._workspace_id = workspace_id
        self._llm_translate_callback = llm_translate_callback

    # ------------------------------------------------------------------
    # FILE TYPE DETECTION
    # ------------------------------------------------------------------

    def _detect_file_type(self, file_path: str) -> str:
        """
        Detect file type using both extension and file signature.

        This is important because Google Drive downloads may return:

            /tmp/Last 5 Emails Summary

        instead of:

            /tmp/Last 5 Emails Summary.pdf
        """

        ext = os.path.splitext(file_path)[1].lower()

        if ext in {".pdf", ".docx", ".txt"}:
            return ext[1:]

        try:
            with open(file_path, "rb") as f:
                header = f.read(16)
        except Exception:
            return "unknown"

        # PDF signature
        if header.startswith(b"%PDF"):
            return "pdf"

        # DOCX is a ZIP container
        if header.startswith(b"PK"):
            try:
                import zipfile

                if zipfile.is_zipfile(file_path):
                    with zipfile.ZipFile(file_path) as z:
                        names = set(z.namelist())

                    if "word/document.xml" in names:
                        return "docx"
            except Exception:
                pass

        # Try UTF-8 text
        try:
            with open(
                file_path,
                "r",
                encoding="utf-8",
                errors="strict",
            ) as f:
                f.read(4096)

            return "txt"
        except Exception:
            pass

        # MIME detection as a final fallback
        mime, _ = mimetypes.guess_type(file_path)

        if mime == "application/pdf":
            return "pdf"

        if mime == (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ):
            return "docx"

        if mime and mime.startswith("text/"):
            return "txt"

        return "unknown"

    # ------------------------------------------------------------------
    # TEXT EXTRACTION
    # ------------------------------------------------------------------

    def _read_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF.

        Uses pypdf first.

        If the PDF contains no extractable text, return a useful
        error so the agent can use OCR instead of treating PDF bytes
        as text.
        """

        try:
            import pypdf
        except ImportError as exc:
            raise RuntimeError(
                "PDF support requires the 'pypdf' package. "
                "Install it with: pip install pypdf"
            ) from exc

        try:
            reader = pypdf.PdfReader(file_path)
        except Exception as exc:
            raise RuntimeError(
                f"Could not open PDF: {exc}"
            ) from exc

        pages = []

        for page_number, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""

                if page_text.strip():
                    pages.append(page_text.strip())

            except Exception as exc:
                logger.warning(
                    "Could not extract page %d from PDF: %s",
                    page_number,
                    exc,
                )

        text = "\n\n".join(pages).strip()

        if text:
            return text

        raise RuntimeError(
            "PDF contains no extractable text. "
            "It may be a scanned/image-only PDF. "
            "Use OCR before translation."
        )

    def _read_docx(self, file_path: str) -> str:
        """Extract paragraphs and basic table content from DOCX."""

        try:
            import docx
        except ImportError as exc:
            raise RuntimeError(
                "DOCX support requires the 'python-docx' package. "
                "Install it with: pip install python-docx"
            ) from exc

        try:
            doc = docx.Document(file_path)
        except Exception as exc:
            raise RuntimeError(
                f"Could not open DOCX: {exc}"
            ) from exc

        parts: list[str] = []

        # Paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            if text:
                parts.append(text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = []

                for cell in row.cells:
                    cell_text = cell.text.strip()

                    if cell_text:
                        cells.append(cell_text)

                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts).strip()

    def _read_txt(self, file_path: str) -> str:
        """Read plain text."""

        encodings = [
            "utf-8",
            "utf-8-sig",
            "cp1252",
            "latin-1",
        ]

        last_error: Exception | None = None

        for encoding in encodings:
            try:
                with open(
                    file_path,
                    "r",
                    encoding=encoding,
                ) as f:
                    return f.read()
            except UnicodeDecodeError as exc:
                last_error = exc

        raise RuntimeError(
            f"Could not decode text file: {last_error}"
        )

    def _read_text(self, file_path: str) -> str:
        """
        Read document content.

        IMPORTANT:
        We don't trust the file extension because Google Drive may
        download a PDF without '.pdf'.
        """

        file_type = self._detect_file_type(file_path)

        logger.info(
            "TranslateDocumentTool detected file type: %s (%s)",
            file_type,
            file_path,
        )

        if file_type == "pdf":
            return self._read_pdf(file_path)

        if file_type == "docx":
            return self._read_docx(file_path)

        if file_type == "txt":
            return self._read_txt(file_path)

        raise RuntimeError(
            f"Unsupported or unknown file type for '{file_path}'. "
            "The file extension may be missing and its content could "
            "not be identified as PDF, DOCX, or TXT."
        )

    # ------------------------------------------------------------------
    # TEXT CLEANING
    # ------------------------------------------------------------------

    def _clean_extracted_text(self, text: str) -> str:
        """
        Clean common extraction artifacts.

        This prevents accidental translation of raw PDF structure such as:

            %PDF-1.4
            1 0 obj
            /BaseFont /Helvetica
            endobj
        """

        if not text:
            return ""

        # Remove null characters
        text = text.replace("\x00", "")

        # Detect raw PDF source accidentally passed as text.
        if text.lstrip().startswith("%PDF-"):
            raise RuntimeError(
                "The file was detected as a PDF, but raw PDF binary/source "
                "content was returned instead of extracted document text. "
                "PDF extraction must be used before translation."
            )

        # Remove excessive blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove excessive spaces
        text = re.sub(r"[ \t]{3,}", " ", text)

        return text.strip()

    # ------------------------------------------------------------------
    # TRANSLATION
    # ------------------------------------------------------------------

    def _translate_with_google_cloud(
        self,
        text: str,
        target_lang: str,
        source_lang: str | None,
    ) -> str:
        """
        Translate using Google Cloud Translation.

        This method supports the common google-cloud-translate package.
        """

        try:
            from google.cloud import translate_v2
        except ImportError as exc:
            raise RuntimeError(
                "Google Cloud Translate package is not installed."
            ) from exc

        try:
            client = translate_v2.Client()
        except Exception as exc:
            raise RuntimeError(
                f"Could not create Google Translate client: {exc}"
            ) from exc

        chunks = self._chunk_text(text, max_chars=4000)

        translated_chunks: list[str] = []

        for chunk in chunks:
            kwargs: dict[str, Any] = {
                "target_language": target_lang,
                "format_": "text",
            }

            if source_lang:
                kwargs["source_language"] = source_lang

            result = client.translate(
                chunk,
                **kwargs,
            )

            translated = result.get("translatedText", "")

            if translated:
                translated_chunks.append(translated)

        result_text = "\n\n".join(translated_chunks).strip()

        if not result_text:
            raise RuntimeError(
                "Google Cloud Translate returned empty text."
            )

        return result_text

    def _translate_with_deep_translator(
        self,
        text: str,
        target_lang: str,
        source_lang: str | None,
    ) -> str:
        """Fallback translation using deep-translator."""

        try:
            from deep_translator import GoogleTranslator
        except ImportError as exc:
            raise RuntimeError(
                "deep-translator is not installed."
            ) from exc

        translator = GoogleTranslator(
            source=source_lang or "auto",
            target=target_lang,
        )

        chunks = self._chunk_text(text, max_chars=4000)

        translated_chunks: list[str] = []

        for chunk in chunks:
            translated = translator.translate(chunk)

            if translated:
                translated_chunks.append(translated)

        result_text = "\n\n".join(
            translated_chunks
        ).strip()

        if not result_text:
            raise RuntimeError(
                "deep-translator returned empty text."
            )

        return result_text

    def _translate_with_llm(
        self,
        text: str,
        target_lang: str,
        source_lang: str | None,
    ) -> str:
        """
        LLM translation fallback.

        The callback is intentionally injectable because your project
        already has its own LLM infrastructure.

        Expected callback:

            callback(prompt) -> translated_text

        or:

            callback(
                text=text,
                target_lang=target_lang,
                source_lang=source_lang
            )
        """

        if not self._llm_translate_callback:
            raise RuntimeError(
                "LLM translation callback is not configured."
            )

        target_name = _LANG_NAMES.get(
            target_lang,
            target_lang,
        )

        source_name = (
            _LANG_NAMES.get(source_lang, source_lang)
            if source_lang
            else "the source language"
        )

        chunks = self._chunk_text(
            text,
            max_chars=6000,
        )

        translated_chunks: list[str] = []

        for chunk in chunks:
            prompt = f"""
Translate the following document text from {source_name}
to {target_name}.

Requirements:
- Translate accurately and naturally.
- Preserve the original meaning.
- Do not summarize.
- Do not add explanations.
- Do not remove important information.
- Preserve headings and paragraph structure where possible.
- Return ONLY the translated text.

DOCUMENT TEXT:
{chunk}
""".strip()

            try:
                translated = self._llm_translate_callback(
                    prompt
                )
            except TypeError:
                translated = self._llm_translate_callback(
                    text=chunk,
                    target_lang=target_lang,
                    source_lang=source_lang,
                )

            if translated:
                translated_chunks.append(
                    str(translated).strip()
                )

        result_text = "\n\n".join(
            translated_chunks
        ).strip()

        if not result_text:
            raise RuntimeError(
                "LLM translation returned empty text."
            )

        return result_text

    def _translate_text(
        self,
        text: str,
        target_lang: str,
        source_lang: str | None,
    ) -> str:
        """
        Try available translation providers.

        Priority:

            1. Google Cloud
            2. Project LLM callback
            3. deep-translator

        The method only returns successfully translated text.
        """

        errors: list[str] = []

        # --------------------------------------------------------------
        # 1. Google Cloud Translate
        # --------------------------------------------------------------

        try:
            translated = self._translate_with_google_cloud(
                text=text,
                target_lang=target_lang,
                source_lang=source_lang,
            )

            logger.info(
                "Google Cloud translation successful: %s -> %s",
                source_lang or "auto",
                target_lang,
            )

            return translated

        except Exception as exc:
            errors.append(
                f"Google Cloud: {exc}"
            )

            logger.warning(
                "Google Cloud Translate unavailable/failed: %s",
                exc,
            )

        # --------------------------------------------------------------
        # 2. LLM fallback
        # --------------------------------------------------------------

        try:
            translated = self._translate_with_llm(
                text=text,
                target_lang=target_lang,
                source_lang=source_lang,
            )

            logger.info(
                "LLM translation successful: %s -> %s",
                source_lang or "auto",
                target_lang,
            )

            return translated

        except Exception as exc:
            errors.append(
                f"LLM: {exc}"
            )

            logger.warning(
                "LLM translation unavailable/failed: %s",
                exc,
            )

        # --------------------------------------------------------------
        # 3. deep-translator fallback
        # --------------------------------------------------------------

        try:
            translated = self._translate_with_deep_translator(
                text=text,
                target_lang=target_lang,
                source_lang=source_lang,
            )

            logger.info(
                "deep-translator successful: %s -> %s",
                source_lang or "auto",
                target_lang,
            )

            return translated

        except Exception as exc:
            errors.append(
                f"deep-translator: {exc}"
            )

            logger.warning(
                "deep-translator unavailable/failed: %s",
                exc,
            )

        raise RuntimeError(
            "Translation failed. "
            + " | ".join(errors)
        )

    # ------------------------------------------------------------------
    # TEXT CHUNKING
    # ------------------------------------------------------------------

    def _chunk_text(
        self,
        text: str,
        max_chars: int = 4000,
    ) -> list[str]:
        """
        Split large documents into chunks.

        Prefer paragraph boundaries instead of cutting sentences
        in the middle.
        """

        if len(text) <= max_chars:
            return [text]

        paragraphs = text.split("\n\n")

        chunks: list[str] = []
        current = ""

        for paragraph in paragraphs:
            paragraph = paragraph.strip()

            if not paragraph:
                continue

            if len(current) + len(paragraph) + 2 <= max_chars:
                current += (
                    ("\n\n" if current else "")
                    + paragraph
                )
                continue

            if current:
                chunks.append(current)

            # Very large paragraph
            if len(paragraph) > max_chars:
                for i in range(
                    0,
                    len(paragraph),
                    max_chars,
                ):
                    chunks.append(
                        paragraph[
                            i:i + max_chars
                        ]
                    )

                current = ""
            else:
                current = paragraph

        if current:
            chunks.append(current)

        return chunks

    # ------------------------------------------------------------------
    # OUTPUT
    # ------------------------------------------------------------------

    def _save_output(
        self,
        text: str,
        file_path: str,
        target_lang: str,
        fmt: str,
    ) -> str:
        """
        Save translated text.

        Default output is DOCX.
        """

        original_name = os.path.basename(
            file_path
        )

        base = os.path.splitext(
            original_name
        )[0]

        # Handle extensionless Google Drive filenames
        if not base:
            base = "translated_document"

        safe_lang = re.sub(
            r"[^a-zA-Z0-9_-]",
            "",
            target_lang,
        )

        if not safe_lang:
            safe_lang = "translated"

        output_base = (
            f"{base}_{safe_lang}"
        )

        if fmt.lower() == "docx":

            try:
                import docx
            except ImportError:
                logger.warning(
                    "python-docx not installed; "
                    "falling back to TXT."
                )
                fmt = "txt"

            else:
                out_path = os.path.join(
                    tempfile.gettempdir(),
                    f"{output_base}.docx",
                )

                doc = docx.Document()

                # Title
                doc.add_heading(
                    (
                        "Translated Document - "
                        f"{_LANG_NAMES.get(target_lang, target_lang)}"
                    ),
                    level=1,
                )

                # Add translated paragraphs
                for paragraph in text.split("\n"):
                    stripped = paragraph.strip()

                    if stripped:
                        doc.add_paragraph(
                            stripped
                        )

                doc.save(out_path)

                logger.info(
                    "Translated DOCX created: %s",
                    out_path,
                )

                return out_path

        # TXT fallback
        out_path = os.path.join(
            tempfile.gettempdir(),
            f"{output_base}.txt",
        )

        with open(
            out_path,
            "w",
            encoding="utf-8",
        ) as f:
            f.write(text)

        logger.info(
            "Translated TXT created: %s",
            out_path,
        )

        return out_path

    # ------------------------------------------------------------------
    # RUN
    # ------------------------------------------------------------------

    def run(
        self,
        input_str: str,
    ) -> str:
        """
        Main tool entry point.
        """

        # --------------------------------------------------------------
        # Parse input
        # --------------------------------------------------------------

        try:
            data = (
                json.loads(input_str)
                if isinstance(input_str, str)
                else input_str
            )

        except (
            json.JSONDecodeError,
            TypeError,
        ):
            return json.dumps(
                {
                    "status": "failed",
                    "error": "Invalid input. Expected JSON.",
                }
            )

        if not isinstance(data, dict):
            return json.dumps(
                {
                    "status": "failed",
                    "error": "Input must be a JSON object.",
                }
            )

        file_path = str(
            data.get("file_path", "")
        ).strip()

        target_lang = str(
            data.get("target_lang", "")
        ).strip().lower()

        source_lang = data.get(
            "source_lang"
        )

        if source_lang:
            source_lang = str(
                source_lang
            ).strip().lower()

        output_format = str(
            data.get(
                "output_format",
                "docx",
            )
        ).strip().lower()

        # --------------------------------------------------------------
        # Validate
        # --------------------------------------------------------------

        if not file_path:
            return json.dumps(
                {
                    "status": "failed",
                    "error": "'file_path' is required.",
                }
            )

        if not target_lang:
            return json.dumps(
                {
                    "status": "failed",
                    "error": (
                        "'target_lang' is required. "
                        f"Options: {list(_LANG_NAMES.keys())}"
                    ),
                }
            )

        if target_lang not in _LANG_NAMES:
            return json.dumps(
                {
                    "status": "failed",
                    "error": (
                        f"Unsupported target language "
                        f"'{target_lang}'. "
                        f"Supported: {list(_LANG_NAMES.keys())}"
                    ),
                }
            )

        if source_lang and source_lang not in _LANG_NAMES:
            return json.dumps(
                {
                    "status": "failed",
                    "error": (
                        f"Unsupported source language "
                        f"'{source_lang}'. "
                        f"Supported: {list(_LANG_NAMES.keys())}"
                    ),
                }
            )

        if output_format not in {
            "docx",
            "txt",
        }:
            return json.dumps(
                {
                    "status": "failed",
                    "error": (
                        "'output_format' must be "
                        "'docx' or 'txt'."
                    ),
                }
            )

        if not os.path.exists(file_path):
            return json.dumps(
                {
                    "status": "failed",
                    "error": (
                        f"File not found: '{file_path}'"
                    ),
                }
            )

        if not os.path.isfile(file_path):
            return json.dumps(
                {
                    "status": "failed",
                    "error": (
                        f"Path is not a file: '{file_path}'"
                    ),
                }
            )

        # --------------------------------------------------------------
        # Process
        # --------------------------------------------------------------

        try:

            logger.info(
                "TranslateDocumentTool started: "
                "file=%s target=%s source=%s",
                file_path,
                target_lang,
                source_lang or "auto",
            )

            # Detect type first
            detected_type = self._detect_file_type(
                file_path
            )

            logger.info(
                "Detected document type: %s",
                detected_type,
            )

            # Extract
            text = self._read_text(
                file_path
            )

            # Clean
            text = self._clean_extracted_text(
                text
            )

            if not text.strip():
                return json.dumps(
                    {
                        "status": "failed",
                        "error": (
                            "Could not extract text "
                            "from the document."
                        ),
                    },
                    ensure_ascii=False,
                )

            logger.info(
                "Extracted %d characters from %s",
                len(text),
                file_path,
            )

            # Translate
            translated = self._translate_text(
                text=text,
                target_lang=target_lang,
                source_lang=source_lang,
            )

            if not translated.strip():
                return json.dumps(
                    {
                        "status": "failed",
                        "error": (
                            "Translation returned empty content."
                        ),
                    },
                    ensure_ascii=False,
                )

            # Save
            out_path = self._save_output(
                text=translated,
                file_path=file_path,
                target_lang=target_lang,
                fmt=output_format,
            )

            word_count = len(
                translated.split()
            )

            # ----------------------------------------------------------
            # Breadcrumb
            # ----------------------------------------------------------

            try:
                import time

                with open(
                    "/tmp/.krypsos_last_translation.json",
                    "w",
                    encoding="utf-8",
                ) as f:
                    json.dump(
                        {
                            "file_path": out_path,
                            "source_file": file_path,
                            "source_lang": source_lang,
                            "target_lang": target_lang,
                            "detected_type": detected_type,
                            "ts": time.time(),
                        },
                        f,
                    )

            except Exception:
                logger.debug(
                    "Could not write translation breadcrumb.",
                    exc_info=True,
                )

            # ----------------------------------------------------------
            # Success
            # ----------------------------------------------------------

            logger.info(
                "TranslateDocumentTool completed: "
                "%s -> %s, words=%d, output=%s",
                file_path,
                target_lang,
                word_count,
                out_path,
            )

            return json.dumps(
                {
                    "status": "translated",
                    "file_path": out_path,
                    "filename": os.path.basename(
                        out_path
                    ),
                    "source_file": file_path,
                    "source_lang": (
                        source_lang or "auto"
                    ),
                    "target_lang": target_lang,
                    "target_name": _LANG_NAMES.get(
                        target_lang,
                        target_lang,
                    ),
                    "detected_file_type": detected_type,
                    "word_count": word_count,
                    "note": (
                        "This file is the final translated document. "
                        "Pass file_path to upload_to_drive if the user "
                        "asked to save it to Google Drive."
                    ),
                },
                ensure_ascii=False,
            )

        except Exception as exc:

            logger.exception(
                "TranslateDocumentTool failed"
            )

            return json.dumps(
                {
                    "status": "failed",
                    "error": str(exc),
                    "file_path": file_path,
                    "target_lang": target_lang,
                },
                ensure_ascii=False,
            )

    # ------------------------------------------------------------------
    # SCHEMA
    # ------------------------------------------------------------------

    def to_schema(self) -> dict:
        return {
            "type": "function",
            "function": {
                "name": self.name,
                "description": self.description,
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": (
                                "Path to PDF/DOCX/TXT file. "
                                "The file may have no extension "
                                "if downloaded from Google Drive."
                            ),
                        },
                        "target_lang": {
                            "type": "string",
                            "enum": list(
                                _LANG_NAMES.keys()
                            ),
                            "description": (
                                "Target language code. "
                                "ta=Tamil, hi=Hindi, "
                                "ml=Malayalam, te=Telugu, "
                                "kn=Kannada, fr=French, "
                                "de=German, es=Spanish, "
                                "ar=Arabic, zh=Chinese, "
                                "ja=Japanese, pt=Portuguese, "
                                "ru=Russian."
                            ),
                        },
                        "source_lang": {
                            "type": "string",
                            "description": (
                                "Source language code. "
                                "Optional; automatic detection "
                                "is used when omitted."
                            ),
                        },
                        "output_format": {
                            "type": "string",
                            "enum": [
                                "docx",
                                "txt",
                            ],
                            "description": (
                                "Output format. Defaults to docx."
                            ),
                        },
                    },
                    "required": [
                        "file_path",
                        "target_lang",
                    ],
                },
            },
        }