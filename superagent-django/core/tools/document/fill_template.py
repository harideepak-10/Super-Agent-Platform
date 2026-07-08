"""
FillTemplateTool — fill a Word (.docx) template with dynamic data.

Zone: GREEN — runs automatically, no human approval required.
"""
from __future__ import annotations
import json
import logging
import os
import shutil
import tempfile
from typing import Any
from core.tools.base_tool import BaseTool, ToolZone

logger = logging.getLogger(__name__)


class FillTemplateTool(BaseTool):
    """Fill a Word (.docx) template with dynamic data using placeholder replacement.

    Placeholders in the template use the format {{FIELD_NAME}} or {{field_name}}.

    Input::

        {
            "template_path": "/tmp/invoice_template.docx",
            "data": {
                "client_name":    "Arun Kumar",
                "invoice_number": "INV-2026-007",
                "amount":         "₹15,000",
                "due_date":       "2026-07-31",
                "company_name":   "KRYPSOS Tech"
            },
            "output_filename": "Invoice_Arun_Kumar.docx"    # optional
        }

    Returns::

        {
            "status":    "filled",
            "file_path": "/tmp/Invoice_Arun_Kumar.docx",
            "filename":  "Invoice_Arun_Kumar.docx",
            "fields_filled": ["client_name", "invoice_number", "amount", "due_date"]
        }
    """

    name: str = "fill_template"
    description: str = (
        "Fill a Word (.docx) template with dynamic data. GREEN — runs automatically. "
        "Template must use {{FIELD_NAME}} placeholders. "
        "Input JSON: {\"template_path\": \"/tmp/invoice.docx\", "
        "\"data\": {\"client_name\": \"Arun\", \"amount\": \"₹15000\"}, "
        "\"output_filename\": \"Invoice_Arun.docx\"}. "
        "Returns file_path — pass to upload_to_drive."
    )
    zone: ToolZone = ToolZone.GREEN

    def __init__(self, workspace_id: str | None = None) -> None:
        self._workspace_id = workspace_id

    def _replace_in_paragraph(self, paragraph, data: dict) -> list:
        """Replace placeholders in a paragraph's runs, return list of filled keys."""
        filled = []
        full_text = "".join(run.text for run in paragraph.runs)

        for key, value in data.items():
            placeholder = "{{" + key + "}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
                filled.append(key)

        if filled:
            # Rebuild runs — put everything in the first run, clear the rest
            for i, run in enumerate(paragraph.runs):
                if i == 0:
                    run.text = full_text
                else:
                    run.text = ""

        return filled

    def run(self, input_str: str) -> str:
        try:
            data_in = json.loads(input_str) if isinstance(input_str, str) else input_str
        except (json.JSONDecodeError, TypeError):
            return json.dumps({"error": "Invalid input."})

        template_path   = data_in.get("template_path", "")
        fill_data       = data_in.get("data", {})
        output_filename = data_in.get("output_filename", "")

        if not template_path:
            return json.dumps({"error": "'template_path' is required."})
        if not os.path.exists(template_path):
            return json.dumps({"error": f"Template not found: '{template_path}'"})
        if not fill_data:
            return json.dumps({"error": "'data' dict is required."})

        try:
            import docx

            # Work on a copy
            tmp_path = os.path.join(tempfile.gettempdir(),
                                    output_filename or os.path.basename(template_path))
            shutil.copy2(template_path, tmp_path)

            doc    = docx.Document(tmp_path)
            filled = []

            # Replace in body paragraphs
            for para in doc.paragraphs:
                filled.extend(self._replace_in_paragraph(para, fill_data))

            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            filled.extend(self._replace_in_paragraph(para, fill_data))

            # Replace in headers and footers
            for section in doc.sections:
                for para in section.header.paragraphs:
                    filled.extend(self._replace_in_paragraph(para, fill_data))
                for para in section.footer.paragraphs:
                    filled.extend(self._replace_in_paragraph(para, fill_data))

            doc.save(tmp_path)
            filled_unique = list(set(filled))

            logger.info("FillTemplateTool: filled %d fields in %s", len(filled_unique), tmp_path)
            return json.dumps({
                "status":       "filled",
                "file_path":    tmp_path,
                "filename":     os.path.basename(tmp_path),
                "fields_filled": filled_unique,
                "note":         "Pass file_path to upload_to_drive to save to Google Drive.",
            }, ensure_ascii=False)

        except ImportError:
            return json.dumps({"error": "python-docx not installed. Run: pip install python-docx"})
        except Exception as exc:
            logger.exception("FillTemplateTool failed")
            return json.dumps({"error": str(exc)})

    def to_schema(self) -> dict:
        return {"type": "function", "function": {
            "name": self.name, "description": self.description,
            "parameters": {"type": "object", "properties": {
                "template_path":   {"type": "string", "description": "Path to .docx template file"},
                "data":            {"type": "object", "description": "Key-value pairs to fill in (keys match {{FIELD}} placeholders)"},
                "output_filename": {"type": "string", "description": "Output filename (optional)"},
            }, "required": ["template_path", "data"]},
        }}
