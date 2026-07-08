"""
ExtractDataFromAttachmentTool — extract structured data (tables, amounts, dates) from an attachment.

Zone: GREEN — runs automatically, no human approval required.

Reads the file content then extracts structured data using regex patterns.
Works best on PDFs and CSVs with tabular data.
"""
from __future__ import annotations
import csv
import json
import logging
import os
import re
from core.tools.base_tool import BaseTool, ToolZone

logger = logging.getLogger(__name__)

_AMOUNT_RE  = re.compile(r'(?:₹|Rs\.?|INR|USD|\$|€|£)\s*([\d,]+(?:\.\d{1,2})?)', re.IGNORECASE)
_DATE_RE    = re.compile(
    r'\b(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4}|\d{4}[\/\-]\d{2}[\/\-]\d{2}|'
    r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+\d{1,2},?\s+\d{4})\b',
    re.IGNORECASE,
)
_EMAIL_RE   = re.compile(r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b')
_PHONE_RE   = re.compile(r'(?:\+91|0)?[\s\-]?[6-9]\d{9}|\b\d{3}[\s\-]\d{3}[\s\-]\d{4}\b')


class ExtractDataFromAttachmentTool(BaseTool):
    """Extract structured data (tables, amounts, dates, contacts) from an attachment.

    Input::

        {
            "file_path": "/tmp/krypsos_docs/report.pdf"
        }

    Returns::

        {
            "filename":   "report.pdf",
            "amounts":    [{"raw": "₹45,000", "numeric": 45000.0}],
            "dates":      ["July 1, 2026", "July 15, 2026"],
            "emails":     ["vendor@example.com"],
            "phones":     ["9876543210"],
            "table_data": [["Name", "Amount", "Date"], ["Invoice #1", "₹45,000", "July 1"]],
            "summary":    "Found 2 amounts, 2 dates, 1 email address."
        }
    """

    name: str = "extract_data_from_attachment"
    description: str = (
        "Extract structured data (amounts, dates, tables, contacts) from a downloaded attachment. "
        "Input JSON: {\"file_path\": \"...\"}. "
        "Get file_path from download_attachment. "
        "Returns amounts, dates, email addresses, phone numbers, and table rows found in the file."
    )
    zone: ToolZone = ToolZone.GREEN

    def run(self, input_str: str) -> str:
        try:
            data = json.loads(input_str) if isinstance(input_str, str) else input_str
        except (json.JSONDecodeError, TypeError):
            return json.dumps({"error": "Invalid input."})

        file_path = data.get("file_path", "")
        if not file_path:
            return json.dumps({"error": "'file_path' is required."})
        if not os.path.exists(file_path):
            return json.dumps({"error": f"File not found: {file_path}"})

        filename = os.path.basename(file_path)
        ext      = os.path.splitext(filename)[1].lower()

        try:
            text, table_data = self._read_file(file_path, ext)
        except Exception as exc:
            logger.exception("ExtractDataFromAttachmentTool: read failed")
            return json.dumps({"error": f"Could not read file: {exc}"})

        amounts = self._extract_amounts(text)
        dates   = list(dict.fromkeys(_DATE_RE.findall(text)))
        emails  = list(dict.fromkeys(_EMAIL_RE.findall(text)))
        phones  = list(dict.fromkeys(_PHONE_RE.findall(text)))

        parts = []
        if amounts:  parts.append(f"{len(amounts)} amount(s)")
        if dates:    parts.append(f"{len(dates)} date(s)")
        if emails:   parts.append(f"{len(emails)} email address(es)")
        if phones:   parts.append(f"{len(phones)} phone number(s)")
        if table_data: parts.append(f"{len(table_data)-1} data row(s)")
        summary = "Found " + ", ".join(parts) + "." if parts else "No structured data found."

        return json.dumps({
            "filename":   filename,
            "amounts":    amounts,
            "dates":      dates,
            "emails":     emails,
            "phones":     phones,
            "table_data": table_data,
            "summary":    summary,
        }, ensure_ascii=False)

    @staticmethod
    def _extract_amounts(text: str) -> list:
        results = []
        for m in _AMOUNT_RE.finditer(text):
            raw = m.group(0).strip()
            try:
                numeric = float(m.group(1).replace(",", ""))
            except ValueError:
                numeric = None
            results.append({"raw": raw, "numeric": numeric})
        return results

    @staticmethod
    def _read_file(file_path: str, ext: str):
        """Return (text, table_data). table_data is list of rows for CSVs."""
        if ext == ".csv":
            rows = []
            text_lines = []
            with open(file_path, "r", encoding="utf-8-sig", errors="replace") as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append(row)
                    text_lines.append(", ".join(row))
            return "\n".join(text_lines), rows

        elif ext == ".pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                text   = "\n".join(p.extract_text() or "" for p in reader.pages)
            except ImportError:
                raise ImportError("pypdf is not installed. Run: pip install pypdf")
            return text, []

        elif ext == ".docx":
            from docx import Document
            doc  = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            # Extract tables from docx
            table_data = []
            for table in doc.tables:
                for row in table.rows:
                    table_data.append([cell.text.strip() for cell in row.cells])
            return text, table_data

        else:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            return text, []

    def to_schema(self) -> dict:
        return {"type": "function", "function": {
            "name": self.name, "description": self.description,
            "parameters": {"type": "object",
                "properties": {"file_path": {"type": "string"}},
                "required": ["file_path"],
            },
        }}
